import logging
import time

from bs4 import BeautifulSoup
from docx import Document
from selenium import webdriver
from selenium.webdriver.chrome.options import Options
from selenium.webdriver.common.by import By
from selenium.webdriver.support import expected_conditions as EC
from selenium.webdriver.support.ui import Select
from selenium.webdriver.support.ui import WebDriverWait

logging.basicConfig(level=logging.INFO)


class MisSiir:
    def __init__(self):
        self.url = 'https://zeusr.sii.cl//AUT2000/InicioAutenticacion/IngresoRutClave.html?https://misiir.sii.cl/cgi_misii/siihome.cgi'
        # self.rut = '76766405-2'
        # self.tax_code = 'GTSPA76'
        self.rut = '76113362-4'
        self.tax_code = 'taleb18'

    @staticmethod
    def config_driver() -> webdriver.Chrome:
        user_agent = "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/58.0.3029.110 Safari/537.3"

        chrome_options = Options()
        # chrome_options.add_argument('--headless')
        chrome_options.add_argument("--start-maximized")
        chrome_options.add_argument("--log-level=3")
        # chrome_options.add_argument(f'user-agent={user_agent}')
        return webdriver.Chrome(options=chrome_options)

    @staticmethod
    def _wait_until_text_visible(driver: webdriver.Chrome, xpath: str, txt: str):
        try:
            WebDriverWait(driver, 20).until(EC.presence_of_element_located((By.XPATH, xpath)))
            if driver.find_element(By.XPATH, xpath).text == txt:
                return
        except Exception as e:
            pass

    def _login(self, driver: webdriver.Chrome):
        driver.get(self.url)
        driver.find_element(By.XPATH, "//input[@id='rutcntr']").send_keys(self.rut)
        driver.find_element(By.XPATH, "//input[@id='clave']").send_keys(self.tax_code)
        driver.find_element(By.XPATH, "//button[@id='bt_ingresar']").click()
        time.sleep(5)
        logging.info('--> logic success.')

    @staticmethod
    def _misii_doc(driver: webdriver.Chrome, doc: Document):
        try:
            name = driver.find_element(By.XPATH, "//p[@id='nameCntr']").text
            rut = driver.find_element(By.XPATH, "//p[@id='rutCntr']").text
            home = driver.find_element(By.XPATH, "//p[@id='domiCntr']").text
            email = driver.find_element(By.XPATH, "//p[@id='mailCntrNoti']").text

            doc.add_heading('MISII', 1)
            doc.add_paragraph(f'Nombre o razón social: {name}')
            doc.add_paragraph(f'RUT: {rut}')
            doc.add_paragraph(f'Domicilio: {home}')
            doc.add_paragraph(f'Correo para notificaciones: {email}')
            logging.info('--> MISII extracted.')
        except Exception as e:
            pass

    @staticmethod
    def _inicio_de_actividades_y_termino_de_giro(driver: webdriver.Chrome, doc: Document):
        try:
            click_ele = driver.find_element(By.XPATH, "//div[@id='headingTwo']//div[@class='ic_arrow']")
            driver.execute_script("arguments[0].click();", click_ele)
            table_xpath = "//div[@id='collapse2Cntrb']//div[@id='no-more-tables']//tbody"
            WebDriverWait(driver, 20).until(EC.presence_of_element_located((By.XPATH, table_xpath)))

            html = driver.find_element(By.XPATH, table_xpath).get_attribute('innerHTML')
            soup = BeautifulSoup(html, 'html.parser')
            td_elements = soup.find_all('td')
            td_values = [td.get_text(strip=True) for td in td_elements]
            sub_lists = [td_values[i:i + 3] for i in range(0, len(td_values), 3)]
            doc.add_heading('Inicio de actividades y término de giro', 1)
            table = doc.add_table(rows=len(sub_lists) + 1, cols=3)
            table.style = 'Table Grid'
            table.cell(0, 0).text = 'Fecha constitución'
            table.cell(0, 1).text = 'Inicio actividades'
            table.cell(0, 2).text = 'Término giro'
            for ind, line in enumerate(sub_lists):
                table.cell(ind + 1, 0).text = line[0]
                table.cell(ind + 1, 1).text = line[1]
                table.cell(ind + 1, 2).text = line[2]
            logging.info('--> Inicio de actividades y término de giro extracted.')

        except Exception as e:
            pass

    @staticmethod
    def _representantes_legales_vigentes(driver: webdriver.Chrome, doc: Document):
        try:
            click_ele = driver.find_element(By.XPATH, "//div[@id='headingConsultas']//div[@class='ic_arrow']")
            driver.execute_script("arguments[0].click();", click_ele)
            table_xpath = "//div[@id='represVig']//div[@id='no-more-tables']//tbody"
            WebDriverWait(driver, 20).until(EC.presence_of_element_located((By.XPATH, table_xpath)))
            html = driver.find_element(By.XPATH, table_xpath).get_attribute('innerHTML')
            soup = BeautifulSoup(html, 'html.parser')
            td_elements = soup.find_all('td')
            td_values = [td.get_text(strip=True) for td in td_elements]
            sub_lists = [td_values[i:i + 3] for i in range(0, len(td_values), 3)]
            doc.add_heading('Representantes legales vigentes', 1)
            table = doc.add_table(rows=len(sub_lists) + 1, cols=3)
            table.style = 'Table Grid'
            table.cell(0, 0).text = 'Nombre'
            table.cell(0, 1).text = 'Rut'
            table.cell(0, 2).text = 'A partir de'
            for ind, line in enumerate(sub_lists):
                table.cell(ind + 1, 0).text = line[0]
                table.cell(ind + 1, 1).text = line[1]
                table.cell(ind + 1, 2).text = line[2]
            logging.info('--> Representantes legales vigentes extracted.')
        except Exception as e:
            pass

    @staticmethod
    def _socios_y_capital(driver: webdriver.Chrome, doc: Document):
        try:
            click_ele = driver.find_element(By.XPATH, "//div[@id='headingActualizacion']")
            driver.execute_script("arguments[0].click();", click_ele)
            table_xpath = "//div[@id='divSociosNew']//div[@id='no-more-tables']//tbody"
            WebDriverWait(driver, 20).until(EC.presence_of_element_located((By.XPATH, table_xpath)))
            html = driver.find_element(By.XPATH, table_xpath).get_attribute('innerHTML')
            soup = BeautifulSoup(html, 'html.parser')
            td_elements = soup.find_all('td')
            td_values = [td.get_text(strip=True) for td in td_elements]
            sub_lists = [td_values[i:i + 8] for i in range(0, len(td_values), 8)]
            doc.add_heading('Socios y Capital', 1)
            table = doc.add_table(rows=len(sub_lists) + 1, cols=7)
            table.style = 'Table Grid'
            table.cell(0, 0).text = 'Nombre'
            table.cell(0, 1).text = 'Rut'
            table.cell(0, 2).text = 'Capital enterado$'
            table.cell(0, 3).text = 'Capital por enterar $'
            table.cell(0, 4).text = '% Capital'
            table.cell(0, 5).text = '% Utilidades'
            table.cell(0, 6).text = 'Fecha de incorporación'
            for ind, line in enumerate(sub_lists):
                table.cell(ind + 1, 0).text = line[0]
                table.cell(ind + 1, 1).text = line[1]
                table.cell(ind + 1, 2).text = line[2]
                table.cell(ind + 1, 3).text = line[3]
                table.cell(ind + 1, 4).text = line[4]
                table.cell(ind + 1, 5).text = line[5]
                table.cell(ind + 1, 6).text = line[6]
            logging.info('--> Socios y Capital extracted.')
        except Exception as e:
            pass

    @staticmethod
    def _actividades_economicas(driver: webdriver.Chrome, doc: Document):
        try:
            click_ele = driver.find_element(By.XPATH, "//div[@id='headingP10']//div[@class='ic_arrow']")
            driver.execute_script("arguments[0].click();", click_ele)
            table_xpath = "//table[@id='tblIdGiros']//tbody"
            WebDriverWait(driver, 20).until(EC.presence_of_element_located((By.XPATH, table_xpath)))
            html = driver.find_element(By.XPATH, table_xpath).get_attribute('innerHTML')
            soup = BeautifulSoup(html, 'html.parser')
            td_elements = soup.find_all('td')
            td_values = [td.get_text(strip=True) for td in td_elements]
            sub_lists = [td_values[i:i + 5] for i in range(0, len(td_values), 5)]
            doc.add_heading('Actividades económicas', 1)
            table = doc.add_table(rows=len(sub_lists) + 1, cols=5)
            table.style = 'Table Grid'
            table.cell(0, 0).text = 'Actividad'
            table.cell(0, 1).text = 'Código'
            table.cell(0, 2).text = 'Categoría tributaria'
            table.cell(0, 3).text = 'Afecta IVA'
            table.cell(0, 4).text = 'A partir de'
            for ind, line in enumerate(sub_lists):
                table.cell(ind + 1, 0).text = line[0]
                table.cell(ind + 1, 1).text = line[1]
                table.cell(ind + 1, 2).text = line[2]
                table.cell(ind + 1, 3).text = line[3]
                table.cell(ind + 1, 4).text = line[4]

            logging.info('--> Actividades económicas extracted.')
        except Exception as e:
            pass

    @staticmethod
    def _sociedades_a_las_que_pertenece_el_contribuyente(driver: webdriver.Chrome, doc: Document):
        try:
            click_ele = driver.find_element(By.XPATH, "//div[@id='headingP11']//div[@class='box_arrow']")
            driver.execute_script("arguments[0].click();", click_ele)
            table_xpath = "//table[@id='idTableMiSoc']//tbody"
            WebDriverWait(driver, 20).until(EC.presence_of_element_located((By.XPATH, table_xpath)))
            html = driver.find_element(By.XPATH, table_xpath).get_attribute('innerHTML')
            soup = BeautifulSoup(html, 'html.parser')
            td_elements = soup.find_all('td')
            td_values = [td.get_text(strip=True) for td in td_elements]
            sub_lists = [td_values[i:i + 9] for i in range(0, len(td_values), 9)]
            doc.add_heading('Sociedades a las que pertenece el contribuyente', 1)
            table = doc.add_table(rows=len(sub_lists) + 1, cols=9)
            table.style = 'Table Grid'
            table.cell(0, 0).text = 'Nombre Sociedad o entes'
            table.cell(0, 1).text = 'Rut'
            table.cell(0, 2).text = 'Termino de Giro'
            table.cell(0, 3).text = 'Capital enterado $'
            table.cell(0, 4).text = 'Capital por enterar $'
            table.cell(0, 5).text = 'Fecha por enterar'
            table.cell(0, 6).text = '% Capital'
            table.cell(0, 7).text = '% participación  utilidades'
            table.cell(0, 8).text = 'Fecha de Incorpación'
            for ind, line in enumerate(sub_lists):
                table.cell(ind + 1, 0).text = line[0]
                table.cell(ind + 1, 1).text = line[1]
                table.cell(ind + 1, 2).text = line[2]
                table.cell(ind + 1, 3).text = line[3]
                table.cell(ind + 1, 4).text = line[4]
                table.cell(ind + 1, 5).text = line[5]
                table.cell(ind + 1, 6).text = line[6]
                table.cell(ind + 1, 7).text = line[7]
                table.cell(ind + 1, 8).text = line[8]

            logging.info('--> Sociedades a las que pertenece el contribuyente extracted.')
        except Exception as e:
            pass

    @staticmethod
    def _caracteristicas_del_contribuyente(driver: webdriver.Chrome, doc: Document):
        try:
            click_ele = driver.find_element(By.XPATH, "//div[@id='ctracc_9']//div[@class='box_arrow']")
            driver.execute_script("arguments[0].click();", click_ele)
            table_xpath = "//div[@id='collapse13Cntrb']//div[@id='no-more-tables']//tbody"
            WebDriverWait(driver, 20).until(EC.presence_of_element_located((By.XPATH, table_xpath)))
            html = driver.find_element(By.XPATH, table_xpath).get_attribute('innerHTML')
            soup = BeautifulSoup(html, 'html.parser')
            td_elements = soup.find_all('td')
            td_values = [td.get_text(strip=True) for td in td_elements]
            sub_lists = [td_values[i:i + 2] for i in range(0, len(td_values), 2)]
            doc.add_heading('Características del contribuyente', 1)
            table = doc.add_table(rows=len(sub_lists) + 1, cols=2)
            table.style = 'Table Grid'
            table.cell(0, 0).text = 'Descripción'
            table.cell(0, 1).text = 'A partir de'
            for ind, line in enumerate(sub_lists):
                table.cell(ind + 1, 0).text = line[0]
                table.cell(ind + 1, 1).text = line[1]
            logging.info('--> Características del contribuyente extracted.')
        except Exception as e:
            pass

    @staticmethod
    def _bienes_raices(driver: webdriver.Chrome, doc: Document):
        try:
            click_ele = driver.find_element(By.XPATH, "//div[@id='headingP14']//div[@class='box_arrow']")
            driver.execute_script("arguments[0].click();", click_ele)
            table_xpath = "//table[@id='tablaAvaluaciones']//tbody"
            WebDriverWait(driver, 20).until(EC.presence_of_element_located((By.XPATH, table_xpath)))
            html = driver.find_element(By.XPATH, table_xpath).get_attribute('innerHTML')
            soup = BeautifulSoup(html, 'html.parser')
            td_elements = soup.find_all('td')
            td_values = [td.get_text(strip=True) for td in td_elements]
            sub_lists = [td_values[i:i + 7] for i in range(0, len(td_values), 7)]
            doc.add_heading('Bienes Raíces', 1)
            table = doc.add_table(rows=len(sub_lists) + 1, cols=7)
            table.style = 'Table Grid'
            table.cell(0, 0).text = ' '
            table.cell(0, 1).text = 'Rol'
            table.cell(0, 2).text = 'Comuna'
            table.cell(0, 3).text = 'Dirección'
            table.cell(0, 4).text = 'Destino'
            table.cell(0, 5).text = 'Cuotas Vencidas por pagar'
            table.cell(0, 6).text = 'Cuotas Vigentes por pagar'
            for ind, line in enumerate(sub_lists):
                if len(line) > 4:
                    table.cell(ind + 1, 0).text = line[0]
                    table.cell(ind + 1, 1).text = line[1]
                    table.cell(ind + 1, 2).text = line[2]
                    table.cell(ind + 1, 3).text = line[3]
                    table.cell(ind + 1, 4).text = line[4]
                    table.cell(ind + 1, 5).text = line[5]
                    table.cell(ind + 1, 6).text = line[6]
            logging.info('--> Bienes Raíces extracted.')
        except Exception as e:
            pass

    @staticmethod
    def _anotaciones_vigentes(driver: webdriver.Chrome, doc: Document):
        try:
            click_ele = driver.find_element(By.XPATH, "//div[@id='comunicaSII5']//div[@class='box_arrow']")
            driver.execute_script("arguments[0].click();", click_ele)
            table_xpath = "//table[@id='idtblAnotacionesCntr']//tbody"
            WebDriverWait(driver, 20).until(EC.presence_of_element_located((By.XPATH, table_xpath)))
            html = driver.find_element(By.XPATH, table_xpath).get_attribute('innerHTML')
            soup = BeautifulSoup(html, 'html.parser')
            td_elements = soup.find_all('td')
            td_values = [td.get_text(strip=True) for td in td_elements]
            sub_lists = [td_values[i:i + 5] for i in range(0, len(td_values), 5)]
            doc.add_heading('Notificaciones', 1)
            table = doc.add_table(rows=len(sub_lists) + 1, cols=7)
            table.style = 'Table Grid'
            table.cell(0, 0).text = 'Anotacion'
            table.cell(0, 1).text = 'Cantidad'
            table.cell(0, 2).text = 'Fecha de activacion'
            table.cell(0, 3).text = 'Descripcion y efectos'
            table.cell(0, 4).text = 'Accion Requerida'
            for ind, line in enumerate(sub_lists):
                table.cell(ind + 1, 0).text = line[0]
                table.cell(ind + 1, 1).text = line[1]
                table.cell(ind + 1, 2).text = line[2]
                table.cell(ind + 1, 3).text = line[3]
                table.cell(ind + 1, 4).text = line[4]
            logging.info('--> Anotaciones Vigentes extracted.')
        except Exception as e:
            pass

    @staticmethod
    def _notificaciones(driver: webdriver.Chrome, doc: Document):
        try:
            click_ele = driver.find_element(By.XPATH, "//div[@id='comunica1']//div[@class='box_arrow']")
            driver.execute_script("arguments[0].click();", click_ele)
            table_xpath = "//table[@id='tablaNotifs']//tbody"
            WebDriverWait(driver, 20).until(EC.presence_of_element_located((By.XPATH, table_xpath)))
            html = driver.find_element(By.XPATH, table_xpath).get_attribute('innerHTML')
            soup = BeautifulSoup(html, 'html.parser')
            td_elements = soup.find_all('td')
            td_values = [td.get_text(strip=True) for td in td_elements]
            sub_lists = [td_values[i:i + 5] for i in range(0, len(td_values), 5)]
            doc.add_heading('Notificaciones', 1)
            table = doc.add_table(rows=len(sub_lists) + 1, cols=7)
            table.style = 'Table Grid'
            table.cell(0, 0).text = 'Id'
            table.cell(0, 1).text = 'Fecha'
            table.cell(0, 2).text = 'Descripción '
            table.cell(0, 3).text = 'Tipo de notificación '
            table.cell(0, 4).text = 'Ver detalle'
            for ind, line in enumerate(sub_lists):
                table.cell(ind + 1, 0).text = line[0]
                table.cell(ind + 1, 1).text = line[1]
                table.cell(ind + 1, 2).text = line[2]
                table.cell(ind + 1, 3).text = line[3]
                table.cell(ind + 1, 4).text = line[4]
            logging.info('--> Notificaciones extracted.')
        except Exception as e:
            pass

    def _registro_de_compras_y_compra(self, driver: webdriver.Chrome, doc):
        years = ["2023", "2024"]
        months = ["01", "02", "03", "04", "05", "06", "07", "08", "09", "10", "11", "12"]
        for year in years:
            for month in months:
                try:
                    month_dropdown = driver.find_element(By.ID, "periodoMes")
                    month_select = Select(month_dropdown)
                    month_select.select_by_value(month)
                    time.sleep(1)
                    year_dropdown = driver.find_element(By.XPATH, "//select[@ng-model='periodoAnho']")
                    year_select = Select(year_dropdown)
                    year_select.select_by_value(year)
                    time.sleep(1)
                    driver.find_element(By.XPATH, "//button[normalize-space()='Consultar']").click()
                    table_xpath = "//table[@class='table table-sm ng-scope']//tbody"
                    WebDriverWait(driver, 5).until(EC.presence_of_element_located((By.XPATH, table_xpath)))
                    html = driver.find_elements(By.XPATH, table_xpath)[1].get_attribute('innerHTML')
                    soup = BeautifulSoup(html, 'html.parser')
                    td_elements = soup.find_all('td')
                    td_values = [td.get_text(strip=True) for td in td_elements]
                    sub_lists = [td_values[i:i + 8] for i in range(0, len(td_values), 8)]
                    header = driver.find_element(By.XPATH, "//strong[@class='ng-binding']").text
                    doc.add_heading(header, 1)
                    table = doc.add_table(rows=len(sub_lists) + 1, cols=8)
                    table.style = 'Table Grid'
                    table.cell(0, 0).text = 'Tipo Documento'
                    table.cell(0, 1).text = 'Total Documentos'
                    table.cell(0, 2).text = 'Monto Exento'
                    table.cell(0, 3).text = 'Monto Neto'
                    table.cell(0, 4).text = 'IVA Recuperable'
                    table.cell(0, 5).text = 'IVA Uso Común'
                    table.cell(0, 6).text = 'IVA No Recuperable'
                    table.cell(0, 7).text = 'Monto Total'
                    for ind, line in enumerate(sub_lists):
                        table.cell(ind + 1, 0).text = line[0]
                        table.cell(ind + 1, 1).text = line[1]
                        table.cell(ind + 1, 2).text = line[2]
                        table.cell(ind + 1, 3).text = line[3]
                        table.cell(ind + 1, 4).text = line[4]
                        table.cell(ind + 1, 5).text = line[5]
                        table.cell(ind + 1, 6).text = line[6]
                        table.cell(ind + 1, 7).text = line[7]
                    logging.info(f'--> {header} extracted.')
                except Exception as e:
                    pass

            year_dropdown = driver.find_element(By.XPATH, "//select[@ng-model='periodoAnho']")
            year_select = Select(year_dropdown)
            year_select.select_by_value(year)
            time.sleep(3)

    def _registro_de_compras_y_venta(self, driver: webdriver.Chrome, doc):
        years = ["2023", "2024"]
        months = ["01", "02", "03", "04", "05", "06", "07", "08", "09", "10", "11", "12"]
        for year in years:
            for month in months:
                try:
                    month_dropdown = driver.find_element(By.ID, "periodoMes")
                    month_select = Select(month_dropdown)
                    month_select.select_by_value(month)
                    time.sleep(1)
                    year_dropdown = driver.find_element(By.XPATH, "//select[@ng-model='periodoAnho']")
                    year_select = Select(year_dropdown)
                    year_select.select_by_value(year)
                    time.sleep(1)
                    driver.find_element(By.XPATH, "//button[normalize-space()='Consultar']").click()
                    table_xpath = "//table[@class='table table-sm ng-scope']//tbody"
                    time.sleep(2)
                    driver.find_element(By.XPATH, "//a[@href='#venta/']").click()
                    WebDriverWait(driver, 5).until(EC.presence_of_element_located((By.XPATH, table_xpath)))
                    html = driver.find_elements(By.XPATH, table_xpath)[1].get_attribute('innerHTML')
                    soup = BeautifulSoup(html, 'html.parser')
                    td_elements = soup.find_all('td')
                    td_values = [td.get_text(strip=True) for td in td_elements]
                    sub_lists = [td_values[i:i + 8] for i in range(0, len(td_values), 8)]

                    header = driver.find_element(By.XPATH, "//strong[@class='ng-binding']").text
                    doc.add_heading(header, 1)
                    table = doc.add_table(rows=len(sub_lists) + 1, cols=8)
                    table.style = 'Table Grid'
                    table.cell(0, 0).text = 'Tipo Documento'
                    table.cell(0, 1).text = 'Total Documentos'
                    table.cell(0, 2).text = 'Monto Exento'
                    table.cell(0, 3).text = 'Monto Neto'
                    table.cell(0, 4).text = 'Monto IVA'
                    table.cell(0, 5).text = 'Monto Total'
                    for ind, line in enumerate(sub_lists):
                        table.cell(ind + 1, 0).text = line[0]
                        table.cell(ind + 1, 1).text = line[1]
                        table.cell(ind + 1, 2).text = line[2]
                        table.cell(ind + 1, 3).text = line[3]
                        table.cell(ind + 1, 4).text = line[4]
                        table.cell(ind + 1, 5).text = line[5]
                    logging.info(f'--> {header} extracted.')
                except Exception as e:
                    pass

            year_dropdown = driver.find_element(By.XPATH, "//select[@ng-model='periodoAnho']")
            year_select = Select(year_dropdown)
            year_select.select_by_value(year)
            time.sleep(3)

    def _page1(self, driver: webdriver.Chrome, doc):
        # misii
        self._misii_doc(driver, doc)

        # click on `Datos personales y tributarios` menu and wait until dom change
        click_ele = driver.find_element(By.XPATH, "//a[@id='menu_datos_contribuyente']")
        driver.execute_script("arguments[0].click();", click_ele)
        self._wait_until_text_visible(driver, "//div[@id='box_datos_contribuyente']//h1[@class='title']",
                                      "Datos personales y tributarios")

        # Inicio de actividades y término de giro
        self._inicio_de_actividades_y_termino_de_giro(driver, doc)

        # Representantes legales vigentes
        self._representantes_legales_vigentes(driver, doc)

        # Socios y Capital
        self._socios_y_capital(driver, doc)

        # Actividades económicas
        self._actividades_economicas(driver, doc)

        # Sociedades a las que pertenece el contribuyente
        self._sociedades_a_las_que_pertenece_el_contribuyente(driver, doc)

        # Características del contribuyente
        self._caracteristicas_del_contribuyente(driver, doc)

        # Bienes Raíces
        self._bienes_raices(driver, doc)

        # click on `SII te informa` menu and wait until dom change
        click_ele = driver.find_element(By.XPATH, "//a[@id='menu_comunicados_sii']")
        driver.execute_script("arguments[0].click();", click_ele)
        self._wait_until_text_visible(driver, "//div[@id='box_comunicados_sii']//h1[contains(@class,'title')]",
                                      "SII te informa")

        # Anotaciones Vigentes
        self._anotaciones_vigentes(driver, doc)

        # Notificaciones
        self._notificaciones(driver, doc)

    def _page2(self, driver: webdriver.Chrome, doc):
        driver.get('https://www4.sii.cl/consdcvinternetui/#/index')
        time.sleep(5)

        # REGISTRO DE COMPRAS Y CAMPRA
        self._registro_de_compras_y_compra(driver, doc)

        # REGISTRO DE COMPRAS Y VENTA
        self._registro_de_compras_y_venta(driver, doc)

    def _page3(self, driver: webdriver.Chrome, doc):
        driver.get('https://www4.sii.cl/sifmConsultaInternet/index.html?form=29&dest=cifxx')
        time.sleep(5)
        try:
            f29 = "//a[normalize-space()='F29 (+)']"
            WebDriverWait(driver, 300).until(EC.presence_of_element_located((By.XPATH, f29)))
            driver.find_element(By.XPATH, f29).click()
            time.sleep(5)
            table_css = "body > div:nth-child(13) > div:nth-child(2) > div:nth-child(1) > table:nth-child(1) > tbody:nth-child(1) > tr:nth-child(2) > td:nth-child(1) > table:nth-child(1) > tbody:nth-child(1) > tr:nth-child(1) > td:nth-child(1) > div:nth-child(1) > table:nth-child(5) > tbody:nth-child(1) > tr:nth-child(1) > td:nth-child(2) > table:nth-child(1) > tbody:nth-child(2) > tr:nth-child(1) > td:nth-child(1) > table:nth-child(1) > tbody:nth-child(1) > tr:nth-child(1) > td:nth-child(1) > table:nth-child(1) > tbody:nth-child(2) > tr:nth-child(3) > td:nth-child(1) > table:nth-child(1) > tbody:nth-child(1) > tr:nth-child(2) > td:nth-child(1) > table:nth-child(1) > tbody:nth-child(1) > tr:nth-child(1) > td:nth-child(1) > table:nth-child(1) > tbody:nth-child(2) > tr:nth-child(1) > td:nth-child(2) > table:nth-child(1) > tbody:nth-child(2)"
            html = driver.find_element(By.CSS_SELECTOR, table_css).get_attribute('innerHTML')
            soup = BeautifulSoup(html, 'html.parser')
            td_elements = soup.find_all('td')
            td_values = [td.get_text(strip=True) for td in td_elements]
            filtered_list = [item for item in td_values if item != ""]
            months_in_spanish = ["Enero", "Febrero", "Marzo", "Abril", "Mayo", "Junio", "Julio", "Agosto", "Septiembre",
                                 "Octubre", "Noviembre", "Diciembre"]
            filtered_list = ["ND" if (
                    "De acuerdo a la información en las bases del SII, debes presentar tu declaración de IVA (F29)." in item and item not in months_in_spanish) else item
                             for item in filtered_list]
            filtered_list = ["-" if (
                    "Periodo sin declaración, haz clic para ingresar al formulario." in item and item not in months_in_spanish) else item
                             for item in filtered_list]
            filtered_list = ["Yes" if ("sin" in item and item not in months_in_spanish) else item for item in
                             filtered_list]
            filtered_list = ["No" if ("con" in item and item not in months_in_spanish) else item for item in
                             filtered_list]
            result, current_list = [], []
            for item in filtered_list:
                if item not in ["Yes", "No", "-", "ND"]:
                    if current_list:
                        result.append(current_list)
                    current_list = [item]
                else:
                    current_list.append(item)
            if current_list:
                result.append(current_list)

            max_length = max(len(inner_list) - 1 for inner_list in result)
            for inner_list in result:
                while len(inner_list) - 1 < max_length:
                    inner_list.insert(1, 'No')

            doc.add_heading("CONSULTA INTEGRAL DE FISCALIZACION", 1)
            table = doc.add_table(rows=len(result) + 1, cols=8)
            table.style = 'Table Grid'
            table.cell(0, 0).text = 'Month'
            table.cell(0, 1).text = '2024'
            table.cell(0, 2).text = '2023'
            table.cell(0, 3).text = '2022'
            table.cell(0, 4).text = '2021'
            table.cell(0, 5).text = '2020'
            table.cell(0, 6).text = '2019'
            table.cell(0, 7).text = '2018'

            for ind, line in enumerate(result):
                table.cell(ind + 1, 0).text = line[0]
                table.cell(ind + 1, 1).text = line[1]
                table.cell(ind + 1, 2).text = line[2]
                table.cell(ind + 1, 3).text = line[3]
                table.cell(ind + 1, 4).text = line[4]
                table.cell(ind + 1, 5).text = line[5]
                table.cell(ind + 1, 6).text = line[6]
                table.cell(ind + 1, 7).text = line[7]

            logging.info(f'--> CONSULTA INTEGRAL DE FISCALIZACION extracted.')

        except Exception as e:
            pass

    def _page4(self, driver: webdriver.Chrome, doc):
        year_value = ["2022", "2023", "2024"]
        month_value = ["Enero", "Febrero", "Marzo", "Abril", "Mayo", "Junio", "Julio", "Agosto", "Septiembre",
                       "Octubre", "Noviembre", "Diciembre"]
        given_codes = ['586', '142', '503', '502', '110', '511', '514', '564', '521', '566', '560', '584', '562', '519',
                       '520', '534', '535', '563', '48', '91']

        for year in year_value:
            for month in month_value:
                try:
                    codes = []
                    driver.get('https://www4.sii.cl/rfiInternet/consulta/index.html#rfiSelFormularioPeriodo')
                    driver.refresh()
                    select_box_title_xpath = "//button[normalize-space()='Buscar Datos Ingresados']"
                    WebDriverWait(driver, 30).until(EC.presence_of_element_located((By.XPATH, select_box_title_xpath)))
                    first_dropdown = driver.find_elements(By.XPATH, '//select[@class="gwt-ListBox"]')[0]
                    first_select = Select(first_dropdown)
                    first_select.select_by_value("0")
                    time.sleep(5)
                    second_dropdown = driver.find_elements(By.XPATH, '//select[@class="gwt-ListBox"]')[1]
                    second_select = Select(second_dropdown)
                    second_select.select_by_visible_text(year)
                    time.sleep(3)
                    third_dropdown = driver.find_elements(By.XPATH, '//select[@class="gwt-ListBox"]')[2]
                    third_select = Select(third_dropdown)
                    third_select.select_by_visible_text(month)
                    time.sleep(3)
                    title_ele = driver.find_element(By.XPATH, select_box_title_xpath)
                    driver.execute_script("arguments[0].click();", title_ele)
                    time.sleep(3)

                    result_xpath = "//td[normalize-space()='DECLARACIONES VIGENTES']"
                    WebDriverWait(driver, 30).until(EC.presence_of_element_located((By.XPATH, result_xpath)))

                    links_ele = driver.find_elements(By.XPATH, '//div[@class="gwt-Hyperlink"]//a')
                    if len(links_ele) > 2:
                        driver.find_elements(By.XPATH, '//div[@class="gwt-Hyperlink"]//a')[2].click()
                        time.sleep(5)
                        driver.find_element(By.XPATH, "//button[normalize-space()='Ver Datos']").click()
                        code_table_xpath = "//div[@class='gwt-HTML']//div[3]"
                        WebDriverWait(driver, 30).until(EC.presence_of_element_located((By.XPATH, code_table_xpath)))
                        trs = driver.find_elements(By.XPATH, "//div[@class='gwt-HTML']//div//tr")
                        for tr in trs:
                            tr_html = tr.get_attribute('innerHTML')
                            if tr_html:
                                soup = BeautifulSoup(tr_html, 'html.parser')
                                td_elements = soup.find_all('td')
                                td_values = [td.get_text(strip=True) for td in td_elements]
                                filtered_items = [item for item in td_values if item not in ['+', '-', '=']]
                                if any(code in filtered_items for code in given_codes) and len(filtered_items) == 6:
                                    codes.append(filtered_items)

                        doc.add_heading(f'FORMULARIO 29 - {month} - {year}', 1)
                        table = doc.add_table(rows=len(codes) + 1, cols=6)
                        table.style = 'Table Grid'
                        for ind, line in enumerate(codes):
                            table.cell(ind + 1, 0).text = line[0]
                            table.cell(ind + 1, 1).text = line[1]
                            table.cell(ind + 1, 2).text = line[2]
                            table.cell(ind + 1, 3).text = line[3]
                            table.cell(ind + 1, 4).text = line[4]
                            table.cell(ind + 1, 5).text = line[5]

                        logging.info(f"--> FORMULARIO 29 - {month} - {year}' extracted.")

                except Exception as e:
                    pass

    def _page5(self, driver: webdriver.Chrome, doc):
        year_value = ["2022", "2023", "2024"]

        for year in year_value:
            try:
                driver.get('https://www4.sii.cl/consultaestadof22ui/#!/default')
                dropdown_xpath = "//select[@class='form-control ng-pristine ng-untouched ng-valid ng-not-empty']"
                WebDriverWait(driver, 30).until(EC.presence_of_element_located((By.XPATH, dropdown_xpath)))
                time.sleep(5)
                dropdown = driver.find_element(By.XPATH, dropdown_xpath)
                dropdown_select = Select(dropdown)
                dropdown_select.select_by_visible_text(year)
                time.sleep(3)
                driver.find_element(By.XPATH, "//button[normalize-space()='Consultar']").click()
                time.sleep(3)
                social_path = "//div[@id='SituacionActual']//span"
                WebDriverWait(driver, 30).until(EC.presence_of_element_located((By.XPATH, social_path)))
                social = driver.find_element(By.XPATH, social_path).text
                time.sleep(10)
                ver_formulario_22_ompacto_button = "//button[normalize-space()='Ver Formulario 22 Compacto']"
                WebDriverWait(driver, 30).until(
                    EC.presence_of_element_located((By.XPATH, ver_formulario_22_ompacto_button)))
                driver.find_element(By.XPATH, ver_formulario_22_ompacto_button).click()

                codes = ["1412", "1430", "1729", "36", "1445", "1459", "1484", "1496", "1513", "1564", "91", "87"]
                data = []

                # 1st container
                table_xpath = '//div[@class="container-fluid container-card no-cabecera"]//div[@class="div-table"]'
                WebDriverWait(driver, 30).until(EC.presence_of_element_located((By.XPATH, table_xpath)))
                tables_ele = driver.find_elements(By.XPATH, table_xpath)
                for item in tables_ele:
                    lines = item.text.split(' ')
                    code = lines[0]
                    name = " ".join(lines[1:-1])
                    number = lines[-1]
                    if code in codes:
                        data.append([code, name, number])

                # 2nd container
                table_xpath = '//div[@class="container-fluid container-card container-result"]//div[@class="div-table"]'
                WebDriverWait(driver, 30).until(EC.presence_of_element_located((By.XPATH, table_xpath)))
                tables_ele = driver.find_elements(By.XPATH, table_xpath)
                for item in tables_ele:
                    lines = item.text.split(' ')
                    lines.pop(-1)
                    matching_indices = [index for index, value in enumerate(lines) if value in codes]
                    if matching_indices:
                        code = lines[matching_indices[0]]
                        lines.pop(matching_indices[0])
                        name = " ".join(lines[0:matching_indices[0]])
                        if len(lines) > matching_indices[0]:
                            number = lines[-1]
                        else:
                            number = ''
                        data.append([code, name, number])

                doc.add_heading(f'CONSULTA DE ESTADO DE DECLARACIÓN DE RENTA {year}', 1)
                doc.add_paragraph(f'Situación Renta Actual: {social}')
                doc.add_heading(f"AÑO TRIBUTARIO {year}", 1)
                table = doc.add_table(rows=len(data) + 1, cols=3)
                table.style = 'Table Grid'
                table.cell(0, 0).text = 'Code'
                table.cell(0, 1).text = 'Name'
                table.cell(0, 2).text = 'Number'
                for ind, line in enumerate(data):
                    table.cell(ind + 1, 0).text = line[0]
                    table.cell(ind + 1, 1).text = line[1]
                    table.cell(ind + 1, 2).text = line[2]

                logging.info(f'Ver Formulario 22 Compacto - {year} extracted')

            except Exception as e:
                pass

    def _page6(self, driver: webdriver.Chrome, doc):
        try:
            driver.get('https://www2.sii.cl/pagogiro-ui/home')
            link_xpath = "//a[normalize-space()='Consulta y Pago de Giros']"
            WebDriverWait(driver, 120).until(EC.presence_of_element_located((By.XPATH, link_xpath)))
            click_ele = driver.find_element(By.XPATH, link_xpath)
            driver.execute_script("arguments[0].click();", click_ele)
            table_xpath = "//table[@id='table-giros']"
            WebDriverWait(driver, 120).until(EC.presence_of_element_located((By.XPATH, table_xpath)))
            html = driver.find_element(By.XPATH, table_xpath).get_attribute('innerHTML')
            soup = BeautifulSoup(html, 'html.parser')
            td_elements = soup.find_all('td')
            td_values = [td.get_text(strip=True) for td in td_elements]
            sub_lists = [td_values[i:i + 9] for i in range(0, len(td_values), 9)]
            doc.add_heading('Inicio de actividades y término de giro', 1)
            table = doc.add_table(rows=len(sub_lists) + 1, cols=8)
            table.style = 'Table Grid'
            table.cell(0, 0).text = 'Folio'
            table.cell(0, 1).text = 'Fec. Emisión'
            table.cell(0, 2).text = 'Fec. Vcto'
            table.cell(0, 3).text = 'Negocio Emisor'
            table.cell(0, 4).text = 'Moneda'
            table.cell(0, 5).text = 'Form'
            table.cell(0, 6).text = 'Total a Pagar'
            table.cell(0, 7).text = 'Acciones'
            for ind, line in enumerate(sub_lists):
                table.cell(ind + 1, 0).text = line[1]
                table.cell(ind + 1, 1).text = line[2]
                table.cell(ind + 1, 2).text = line[3]
                table.cell(ind + 1, 3).text = line[4]
                table.cell(ind + 1, 4).text = line[5]
                table.cell(ind + 1, 5).text = line[6]
                table.cell(ind + 1, 6).text = line[7]
                table.cell(ind + 1, 7).text = line[8]
            logging.info('--> Consulta y Pago de Giros extracted.')
        except Exception as e:
            pass

    def _page7(self, driver: webdriver.Chrome, doc):
        years = ["2023", "2024", "2022"]
        for year in years:
            try:
                driver.get('https://www4.sii.cl/perfilamientodjui/#/declaracionJuradaRenta')
                dropdown_xpath = "//select[@name='anioTributario']"
                WebDriverWait(driver, 50).until(EC.presence_of_element_located((By.XPATH, dropdown_xpath)))
                dropdown = driver.find_element(By.XPATH, dropdown_xpath)
                dropdown_select = Select(dropdown)
                dropdown_select.select_by_visible_text(year)
                driver.find_element(By.XPATH, "//button[normalize-space()='Ir']").click()
                time.sleep(10)
                table_xpath = "//table[@st-table='datosMostrarRegimen']"
                WebDriverWait(driver, 30).until(EC.presence_of_element_located((By.XPATH, table_xpath)))
                html = driver.find_element(By.XPATH, table_xpath).get_attribute('innerHTML')
                soup = BeautifulSoup(html, 'html.parser')
                td_elements = soup.find_all('td')
                td_values = [td.get_text(strip=True) for td in td_elements]
                sub_lists = [td_values[i:i + 6] for i in range(0, len(td_values), 6)]
                doc.add_heading(f'Declaraciones por Régimen - {year}', 1)
                table = doc.add_table(rows=len(sub_lists) + 1, cols=4)
                table.style = 'Table Grid'
                table.cell(0, 0).text = 'codigo'
                table.cell(0, 1).text = 'Declaración Jurada'
                table.cell(0, 2).text = 'Fec. Vcto'
                table.cell(0, 3).text = 'Fecha Presentación'
                for ind, line in enumerate(sub_lists):
                    if len(line) > 4:
                        table.cell(ind + 1, 0).text = line[0]
                        table.cell(ind + 1, 1).text = line[1]
                        table.cell(ind + 1, 2).text = line[2]
                        table.cell(ind + 1, 3).text = line[3]
            except Exception as e:
                pass

    def run(self):
        logging.info('----------------- SCRIPT STARTS -------------------')
        driver = self.config_driver()
        # driver.minimize_window()
        doc = Document()
        self._login(driver)
        logging.info('==== PAGE 1 STARTED ====')
        self._page1(driver, doc)
        logging.info('==== PAGE 2 STARTED ====')
        self._page2(driver, doc)
        logging.info('==== PAGE 3 STARTED ====')
        self._page3(driver, doc)
        logging.info('==== PAGE 4 STARTED ====')
        self._page4(driver, doc)
        logging.info('==== PAGE 5 STARTED ====')
        self._page5(driver, doc)
        logging.info('==== PAGE 6 STARTED ====')
        self._page6(driver, doc)
        logging.info('==== PAGE 7 STARTED ====')
        self._page7(driver, doc)
        doc.save(f'{self.rut}.docx')

        logging.info('----------------- SCRIPT ENDS -------------------')


MisSiir().run()

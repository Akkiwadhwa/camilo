from datetime import datetime
import logging
import re
import time
import calendar
import os
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from selenium import webdriver
from selenium.webdriver.chrome.options import Options
from selenium.webdriver.common.by import By
from selenium.webdriver.support import expected_conditions as EC
from selenium.webdriver.support.ui import Select
from selenium.webdriver.support.ui import WebDriverWait
from selenium.common.exceptions import NoSuchElementException
from io import BytesIO

logging.basicConfig(level=logging.INFO)


class MisSiir:
    def __init__(self, output_dir=None):
        self.url = 'https://zeusr.sii.cl//AUT2000/InicioAutenticacion/IngresoRutClave.html?https://misiir.sii.cl/cgi_misii/siihome.cgi'
        # self.rut = '77769810-9''
        # self.tax_code = 'ARK7776'
        self.rut = '76.104.365-K'
        self.tax_code = 'MAITENES7610'
        self.output_dir = output_dir
        self.driver = None

    @staticmethod
    def config_driver() -> webdriver.Chrome:
        user_agent = "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/58.0.3029.110 Safari/537.3"

        chrome_options = Options()
        chrome_options.add_argument('--headless')
        chrome_options.add_argument("--start-maximized")
        chrome_options.add_argument("--log-level=3")
        chrome_options.add_argument(f'user-agent={user_agent}')
        chrome_options.add_argument("--disable-gpu")
        chrome_options.add_argument("--no-sandbox")
        chrome_options.add_argument("--disable-extensions")
        chrome_options.add_argument("--disable-dev-shm-usage")
        return webdriver.Chrome(options=chrome_options)
    
    def _fix_three_col_widths(self, table, wide=Inches(4), narrow=Inches(1)):
        """
        Force 3‑column table:  first & third narrow,  middle widest.
        Call *after* you’ve finished writing the table content.
        """
        if len(table.columns) != 3:
            return                      # silently ignore non‑3‑col tables
        widths = [narrow, wide, narrow]
        for col, w in zip(table.columns, widths):
            for cell in col.cells:
                cell.width = w      

    @staticmethod
    def _wait_until_text_visible(driver: webdriver.Chrome, xpath: str, txt: str):
        try:
            WebDriverWait(driver, 20).until(EC.presence_of_element_located((By.XPATH, xpath)))
            if driver.find_element(By.XPATH, xpath).text == txt:
                return
        except Exception as e:
            pass

    def _login(self, driver: webdriver.Chrome):
        driver.get(self.url)
        driver.find_element(By.XPATH, "//input[@id='rutcntr']").send_keys(self.rut)
        driver.find_element(By.XPATH, "//input[@id='clave']").send_keys(self.tax_code)
        driver.find_element(By.XPATH, "//button[@id='bt_ingresar']").click()
        time.sleep(5)
        logging.info('--> logic success.')

    @staticmethod
    def _misii_doc(driver: webdriver.Chrome, doc: Document):
        try:
            name = driver.find_element(By.XPATH, "//p[@id='nameCntr']").text
            rut = driver.find_element(By.XPATH, "//p[@id='rutCntr']").text
            home = driver.find_element(By.XPATH, "//p[@id='domiCntr']").text
            email = driver.find_element(By.XPATH, "//p[@id='mailCntrNoti']").text

            doc.add_heading('MISII', 1)
            doc.add_paragraph(f'Nombre o razón social: {name}')
            doc.add_paragraph(f'RUT: {rut}')
            doc.add_paragraph(f'Domicilio: {home}')
            doc.add_paragraph(f'Correo para notificaciones: {email}')
            logging.info('--> MISII extracted.')
        except Exception as e:
            pass

    @staticmethod
    def _inicio_de_actividades_y_termino_de_giro(driver: webdriver.Chrome, doc: Document):
        try:
            click_ele = driver.find_element(By.XPATH, "//*[@id=\"headingTwo\"]/h4/a/div/div/div")
            driver.execute_script("arguments[0].click();", click_ele)
            time.sleep(1)
            table_xpath = "//div[@id='collapse2Cntrb']//div[@id='no-more-tables']//tbody"
            WebDriverWait(driver, 20).until(EC.presence_of_element_located((By.XPATH, table_xpath)))

            html = driver.find_element(By.XPATH, table_xpath).get_attribute('innerHTML')
            soup = BeautifulSoup(html, 'html.parser')
            td_elements = soup.find_all('td')
            td_values = [td.get_text(strip=True) for td in td_elements]
            sub_lists = [td_values[i:i + 3] for i in range(0, len(td_values), 3)]
            # table_element = driver.find_element(By.XPATH, "//*[@id=\"no-more-tables\"]/table")
            # screenshot_png = table_element.screenshot_as_png  # bytes
            # doc.add_heading('Inicio de actividades y término de giro', 1)  
            # image_stream = BytesIO(screenshot_png)
            # doc.add_picture(image_stream, width=Inches(6))  # Ajusta el ancho si es necesario



            table = doc.add_table(rows=len(sub_lists) + 1, cols=3)
            table.style = 'Table Grid'
            table.alignment = WD_TABLE_ALIGNMENT.CENTER

            headers = ['Fecha constitución', 'Inicio actividades', 'Término giro']

            # Encabezados
            for col, title in enumerate(headers):
                cell = table.cell(0, col)
                cell.text = title
                cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                run = cell.paragraphs[0].runs[0]
                run.font.bold = True
                run.font.size = Pt(10)
                cell_shading = OxmlElement('w:shd')
                cell_shading.set(qn('w:fill'), "D9D9D9")  # Gris claro
                cell._tc.get_or_add_tcPr().append(cell_shading)

            # Contenido
            for ind, line in enumerate(sub_lists):
                for col, value in enumerate(line):
                    cell = table.cell(ind + 1, col)
                    cell.text = value
                    cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    run = cell.paragraphs[0].runs[0]
                    run.font.size = Pt(9)

            logging.info('--> Inicio de actividades y término de giro extracted.')

        except Exception as e:
            logging.error(f"Error extrayendo tabla: {e}")

    def _representantes_legales_vigentes(self,driver: webdriver.Chrome, doc: Document):
        try:
            click_ele = driver.find_element(By.XPATH, "//div[@id='headingConsultas']//div[@class='ic_arrow']")
            driver.execute_script("arguments[0].click();", click_ele)
            table_xpath = "//div[@id='represVig']//div[@id='no-more-tables']//tbody"
            WebDriverWait(driver, 20).until(EC.presence_of_element_located((By.XPATH, table_xpath)))

            html = driver.find_element(By.XPATH, table_xpath).get_attribute('innerHTML')
            soup = BeautifulSoup(html, 'html.parser')
            td_elements = soup.find_all('td')
            td_values = [td.get_text(strip=True) for td in td_elements]
            sub_lists = [td_values[i:i + 3] for i in range(0, len(td_values), 3)]

            doc.add_heading('Representantes legales vigentes', 1)
            table = doc.add_table(rows=len(sub_lists) + 1, cols=3)
            table.style = 'Table Grid'
            table.alignment = WD_TABLE_ALIGNMENT.CENTER

            headers = ['Nombre', 'RUT', 'A partir de']

            # Encabezados
            for col, title in enumerate(headers):
                cell = table.cell(0, col)
                cell.text = title
                cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                run = cell.paragraphs[0].runs[0]
                run.font.bold = True
                run.font.size = Pt(10)

                # Color de fondo
                cell_shading = OxmlElement('w:shd')
                cell_shading.set(qn('w:fill'), "D9D9D9")  # gris claro
                cell._tc.get_or_add_tcPr().append(cell_shading)

            # Filas con contenido
            for ind, line in enumerate(sub_lists):
                for col, value in enumerate(line):
                    cell = table.cell(ind + 1, col)
                    cell.text = value
                    cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    run = cell.paragraphs[0].runs[0]
                    run.font.size = Pt(9)

            logging.info('--> Representantes legales vigentes extracted.')
            self._fix_three_col_widths(table)


        except Exception as e:
            logging.error(f"Error extrayendo representantes legales: {e}")

      

    @staticmethod
    def _socios_y_capital(driver: webdriver.Chrome, doc: Document):
        try:
            click_ele = driver.find_element(By.XPATH, "//div[@id='headingActualizacion']")
            driver.execute_script("arguments[0].click();", click_ele)
            table_xpath = "//div[@id='divSociosNew']//div[@id='no-more-tables']//tbody"
            WebDriverWait(driver, 20).until(EC.presence_of_element_located((By.XPATH, table_xpath)))

            html = driver.find_element(By.XPATH, table_xpath).get_attribute('innerHTML')
            soup = BeautifulSoup(html, 'html.parser')
            td_elements = soup.find_all('td')
            td_values = [td.get_text(strip=True) for td in td_elements]
            sub_lists = [td_values[i:i + 8] for i in range(0, len(td_values), 8)]

            doc.add_heading('Socios y Capital', 1)
            table = doc.add_table(rows=len(sub_lists) + 1, cols=7)
            table.style = 'Table Grid'
            table.alignment = WD_TABLE_ALIGNMENT.CENTER

            headers = [
                'Nombre', 'RUT', 'Capital enterado $', 'Capital por enterar $',
                '% Capital', '% Utilidades', 'Fecha de incorporación'
            ]

            # Encabezados
            for col, title in enumerate(headers):
                cell = table.cell(0, col)
                cell.text = title
                cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                run = cell.paragraphs[0].runs[0]
                run.font.bold = True
                run.font.size = Pt(10)

                # Fondo gris
                cell_shading = OxmlElement('w:shd')
                cell_shading.set(qn('w:fill'), "D9D9D9")
                cell._tc.get_or_add_tcPr().append(cell_shading)

            # Filas de contenido
            for ind, line in enumerate(sub_lists):
                for col in range(7):  # Se ignora el índice 7 (columna 8), ya que no está en el diseño
                    cell = table.cell(ind + 1, col)
                    cell.text = line[col]
                    cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    run = cell.paragraphs[0].runs[0]
                    run.font.size = Pt(9)

            logging.info('--> Socios y Capital extracted.')

        except Exception as e:
            logging.error(f"Error extrayendo Socios y Capital: {e}")

    @staticmethod
    def _actividades_economicas(driver: webdriver.Chrome, doc: Document):
        try:
            click_ele = driver.find_element(By.XPATH, "//div[@id='headingP10']//div[@class='ic_arrow']")
            driver.execute_script("arguments[0].click();", click_ele)
            table_xpath = "//table[@id='tblIdGiros']//tbody"
            WebDriverWait(driver, 20).until(EC.presence_of_element_located((By.XPATH, table_xpath)))

            html = driver.find_element(By.XPATH, table_xpath).get_attribute('innerHTML')
            soup = BeautifulSoup(html, 'html.parser')
            td_elements = soup.find_all('td')
            td_values = [td.get_text(strip=True) for td in td_elements]
            sub_lists = [td_values[i:i + 5] for i in range(0, len(td_values), 5)]

            doc.add_heading('Actividades económicas', 1)
            table = doc.add_table(rows=len(sub_lists) + 1, cols=5)
            table.style = 'Table Grid'
            table.alignment = WD_TABLE_ALIGNMENT.CENTER

            table.columns[0].width = Inches(2.5)  # Actividad
            table.columns[1].width = Inches(1.0)  # Código
            table.columns[2].width = Inches(1.5)  # Categoría tributaria
            table.columns[3].width = Inches(0.7)  # Afecta IVA
            table.columns[4].width = Inches(1.0)  # A partir de

            headers = ['Actividad', 'Código', 'Categoría tributaria', 'Afecta IVA', 'A partir de']

            # Encabezados
            for col, title in enumerate(headers):
                cell = table.cell(0, col)
                cell.text = title
                cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                run = cell.paragraphs[0].runs[0]
                run.font.bold = True
                run.font.size = Pt(10)

                # Fondo gris
                cell_shading = OxmlElement('w:shd')
                cell_shading.set(qn('w:fill'), "D9D9D9")
                cell._tc.get_or_add_tcPr().append(cell_shading)

            # Filas con contenido
            for ind, line in enumerate(sub_lists):
                for col in range(5):
                    cell = table.cell(ind + 1, col)
                    cell.text = line[col]
                    cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    run = cell.paragraphs[0].runs[0]
                    run.font.size = Pt(9)

            logging.info('--> Actividades económicas extracted.')

        except Exception as e:
            logging.error(f"Error extrayendo actividades económicas: {e}")

    @staticmethod
    def _sociedades_a_las_que_pertenece_el_contribuyente(driver: webdriver.Chrome, doc: Document):
        try:
            click_ele = driver.find_element(By.XPATH, "//div[@id='headingP11']//div[@class='box_arrow']")
            driver.execute_script("arguments[0].click();", click_ele)
            table_xpath = "//table[@id='idTableMiSoc']//tbody"
            WebDriverWait(driver, 20).until(EC.presence_of_element_located((By.XPATH, table_xpath)))

            html = driver.find_element(By.XPATH, table_xpath).get_attribute('innerHTML')
            soup = BeautifulSoup(html, 'html.parser')
            td_elements = soup.find_all('td')
            td_values = [td.get_text(strip=True) for td in td_elements]
            sub_lists = [td_values[i:i + 9] for i in range(0, len(td_values), 9)]

            doc.add_heading('Sociedades a las que pertenece el contribuyente', 1)
            table = doc.add_table(rows=len(sub_lists) + 1, cols=9)
            table.style = 'Table Grid'
            table.alignment = WD_TABLE_ALIGNMENT.CENTER

            headers = [
                'Nombre Sociedad o entes', 'RUT', 'Término de giro',
                'Capital enterado $', 'Capital por enterar $', 'Fecha por enterar',
                '% Capital', '% Participación utilidades', 'Fecha de incorporación'
            ]

            # Encabezados
            for col, title in enumerate(headers):
                cell = table.cell(0, col)
                cell.text = title
                cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                run = cell.paragraphs[0].runs[0]
                run.font.bold = True
                run.font.size = Pt(10)

                cell_shading = OxmlElement('w:shd')
                cell_shading.set(qn('w:fill'), "D9D9D9")  # Gris claro
                cell._tc.get_or_add_tcPr().append(cell_shading)

            # Filas con datos
            for ind, line in enumerate(sub_lists):
                for col in range(9):
                    cell = table.cell(ind + 1, col)
                    cell.text = line[col]
                    cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    run = cell.paragraphs[0].runs[0]
                    run.font.size = Pt(9)

            logging.info('--> Sociedades a las que pertenece el contribuyente extracted.')

        except Exception as e:
            logging.error(f"Error extrayendo sociedades: {e}")

    def _caracteristicas_del_contribuyente(self,driver: webdriver.Chrome, doc: Document):
        try:
            click_ele = driver.find_element(By.XPATH, "//div[@id='ctracc_9']//div[@class='box_arrow']")
            driver.execute_script("arguments[0].click();", click_ele)
            table_xpath = "//div[@id='collapse13Cntrb']//div[@id='no-more-tables']//tbody"
            WebDriverWait(driver, 20).until(EC.presence_of_element_located((By.XPATH, table_xpath)))

            html = driver.find_element(By.XPATH, table_xpath).get_attribute('innerHTML')
            soup = BeautifulSoup(html, 'html.parser')
            td_elements = soup.find_all('td')
            td_values = [td.get_text(strip=True) for td in td_elements]
            sub_lists = [td_values[i:i + 2] for i in range(0, len(td_values), 2)]

            doc.add_heading('Características del contribuyente', 1)
            table = doc.add_table(rows=len(sub_lists) + 1, cols=2)
            table.style = 'Table Grid'
            table.alignment = WD_TABLE_ALIGNMENT.CENTER

            headers = ['Descripción', 'A partir de']

            # Encabezados
            for col, title in enumerate(headers):
                cell = table.cell(0, col)
                cell.text = title
                cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                run = cell.paragraphs[0].runs[0]
                run.font.bold = True
                run.font.size = Pt(10)

                # Fondo gris claro
                cell_shading = OxmlElement('w:shd')
                cell_shading.set(qn('w:fill'), "D9D9D9")
                cell._tc.get_or_add_tcPr().append(cell_shading)

            # Filas de contenido
            for ind, line in enumerate(sub_lists):
                for col in range(2):
                    cell = table.cell(ind + 1, col)
                    cell.text = line[col]
                    cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    run = cell.paragraphs[0].runs[0]
                    run.font.size = Pt(9)

            logging.info('--> Características del contribuyente extracted.')
            self._fix_three_col_widths(table)


        except Exception as e:
            logging.error(f"Error extrayendo características del contribuyente: {e}")

    @staticmethod
    def _bienes_raices(driver: webdriver.Chrome, doc: Document):
        try:
            click_ele = driver.find_element(By.XPATH, "//div[@id='headingP14']//div[@class='box_arrow']")
            driver.execute_script("arguments[0].click();", click_ele)
            table_xpath = "//table[@id='tablaAvaluaciones']//tbody"
            WebDriverWait(driver, 20).until(EC.presence_of_element_located((By.XPATH, table_xpath)))

            html = driver.find_element(By.XPATH, table_xpath).get_attribute('innerHTML')
            soup = BeautifulSoup(html, 'html.parser')
            td_elements = soup.find_all('td')
            td_values = [td.get_text(strip=True) for td in td_elements]
            sub_lists = [td_values[i:i + 7] for i in range(0, len(td_values), 7)]

            doc.add_heading('Bienes Raíces', 1)
            table = doc.add_table(rows=len(sub_lists) + 1, cols=7)
            table.style = 'Table Grid'
            table.alignment = WD_TABLE_ALIGNMENT.CENTER

            headers = [
                'Identificador', 'Rol', 'Comuna', 'Dirección',
                'Destino', 'Cuotas vencidas por pagar', 'Cuotas vigentes por pagar'
            ]

            # Encabezados
            for col, title in enumerate(headers):
                cell = table.cell(0, col)
                cell.text = title
                cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                run = cell.paragraphs[0].runs[0]
                run.font.bold = True
                run.font.size = Pt(10)

                # Fondo gris
                cell_shading = OxmlElement('w:shd')
                cell_shading.set(qn('w:fill'), "D9D9D9")
                cell._tc.get_or_add_tcPr().append(cell_shading)

            # Contenido
            for ind, line in enumerate(sub_lists):
                if len(line) == 7:
                    for col in range(7):
                        cell = table.cell(ind + 1, col)
                        cell.text = line[col]
                        cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                        run = cell.paragraphs[0].runs[0]
                        run.font.size = Pt(9)

            logging.info('--> Bienes Raíces extracted.')

        except Exception as e:
            logging.error(f"Error extrayendo bienes raíces: {e}")

    @staticmethod
    def _anotaciones_vigentes(driver: webdriver.Chrome, doc: Document):
        try:
            click_ele = driver.find_element(By.XPATH, "//div[@id='comunicaSII5']//div[@class='box_arrow']")
            driver.execute_script("arguments[0].click();", click_ele)
            table_xpath = "//table[@id='idtblAnotacionesCntr']//tbody"
            WebDriverWait(driver, 20).until(EC.presence_of_element_located((By.XPATH, table_xpath)))

            html = driver.find_element(By.XPATH, table_xpath).get_attribute('innerHTML')
            soup = BeautifulSoup(html, 'html.parser')
            td_elements = soup.find_all('td')
            td_values = [td.get_text(strip=True) for td in td_elements]
            sub_lists = [td_values[i:i + 5] for i in range(0, len(td_values), 5)]

            doc.add_heading('Notificaciones', 1)
            table = doc.add_table(rows=len(sub_lists) + 1, cols=5)
            table.style = 'Table Grid'
            table.alignment = WD_TABLE_ALIGNMENT.CENTER

            headers = [
                'Anotación', 'Cantidad', 'Fecha de activación',
                'Descripción y efectos', 'Acción requerida'
            ]

            # Encabezados
            for col, title in enumerate(headers):
                cell = table.cell(0, col)
                cell.text = title
                cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                run = cell.paragraphs[0].runs[0]
                run.font.bold = True
                run.font.size = Pt(10)

                # Fondo gris claro
                cell_shading = OxmlElement('w:shd')
                cell_shading.set(qn('w:fill'), "D9D9D9")
                cell._tc.get_or_add_tcPr().append(cell_shading)

            # Filas de contenido
            for ind, line in enumerate(sub_lists):
                for col in range(5):
                    cell = table.cell(ind + 1, col)
                    cell.text = line[col]
                    cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    run = cell.paragraphs[0].runs[0]
                    run.font.size = Pt(9)

            logging.info('--> Anotaciones Vigentes extracted.')

        except Exception as e:
            logging.error(f"Error extrayendo anotaciones vigentes: {e}")

    @staticmethod
    def _notificaciones(driver: webdriver.Chrome, doc: Document):
        try:
            click_ele = driver.find_element(By.XPATH, "//div[@id='comunica1']//div[@class='box_arrow']")
            driver.execute_script("arguments[0].click();", click_ele)
            table_xpath = "//table[@id='tablaNotifs']//tbody"
            WebDriverWait(driver, 20).until(EC.presence_of_element_located((By.XPATH, table_xpath)))

            html = driver.find_element(By.XPATH, table_xpath).get_attribute('innerHTML')
            soup = BeautifulSoup(html, 'html.parser')
            td_elements = soup.find_all('td')
            td_values = [td.get_text(strip=True) for td in td_elements]
            sub_lists = [td_values[i:i + 5] for i in range(0, len(td_values), 5)]

            doc.add_heading('Notificaciones', 1)
            table = doc.add_table(rows=len(sub_lists) + 1, cols=5)
            table.style = 'Table Grid'
            table.alignment = WD_TABLE_ALIGNMENT.CENTER

            headers = [
                'ID', 'Fecha', 'Descripción',
                'Tipo de notificación', 'Ver detalle'
            ]

            # Encabezados
            for col, title in enumerate(headers):
                cell = table.cell(0, col)
                cell.text = title
                cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                run = cell.paragraphs[0].runs[0]
                run.font.bold = True
                run.font.size = Pt(10)

                # Fondo gris claro
                cell_shading = OxmlElement('w:shd')
                cell_shading.set(qn('w:fill'), "D9D9D9")
                cell._tc.get_or_add_tcPr().append(cell_shading)

            # Filas de contenido
            for ind, line in enumerate(sub_lists):
                for col in range(5):
                    cell = table.cell(ind + 1, col)
                    cell.text = line[col]
                    cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    run = cell.paragraphs[0].runs[0]
                    run.font.size = Pt(9)

            logging.info('--> Notificaciones extracted.')

        except Exception as e:
            logging.error(f"Error extrayendo notificaciones: {e}")

    def _registro_de_compras_y_compra(self, driver: webdriver.Chrome, doc):
        years = ["2023", "2024"]
        months = ["01", "02", "03", "04", "05", "06", "07", "08", "09", "10", "11", "12"]
        for year in years:
            for month in months:
                try:
                    month_dropdown = driver.find_element(By.ID, "periodoMes")
                    month_select = Select(month_dropdown)
                    month_select.select_by_value(month)
                    time.sleep(1)
                    year_dropdown = driver.find_element(By.XPATH, "//select[@ng-model='periodoAnho']")
                    year_select = Select(year_dropdown)
                    year_select.select_by_value(year)
                    time.sleep(1)
                    driver.find_element(By.XPATH, "//button[normalize-space()='Consultar']").click()
                    table_xpath = "//table[@class='table table-sm ng-scope']//tbody"
                    WebDriverWait(driver, 5).until(EC.presence_of_element_located((By.XPATH, table_xpath)))
                    html = driver.find_elements(By.XPATH, table_xpath)[1].get_attribute('innerHTML')
                    soup = BeautifulSoup(html, 'html.parser')
                    td_elements = soup.find_all('td')
                    td_values = [td.get_text(strip=True) for td in td_elements]
                    sub_lists = [td_values[i:i + 8] for i in range(0, len(td_values), 8)]
                    header = driver.find_element(By.XPATH, "//strong[@class='ng-binding']").text
                    doc.add_heading(header, 1)
                    table = doc.add_table(rows=len(sub_lists) + 1, cols=8)
                    table.style = 'Table Grid'
                    table.cell(0, 0).text = 'Tipo Documento'
                    table.cell(0, 1).text = 'Total Documentos'
                    table.cell(0, 2).text = 'Monto Exento'
                    table.cell(0, 3).text = 'Monto Neto'
                    table.cell(0, 4).text = 'IVA Recuperable'
                    table.cell(0, 5).text = 'IVA Uso Común'
                    table.cell(0, 6).text = 'IVA No Recuperable'
                    table.cell(0, 7).text = 'Monto Total'
                    for ind, line in enumerate(sub_lists):
                        table.cell(ind + 1, 0).text = line[0]
                        table.cell(ind + 1, 1).text = line[1]
                        table.cell(ind + 1, 2).text = line[2]
                        table.cell(ind + 1, 3).text = line[3]
                        table.cell(ind + 1, 4).text = line[4]
                        table.cell(ind + 1, 5).text = line[5]
                        table.cell(ind + 1, 6).text = line[6]
                        table.cell(ind + 1, 7).text = line[7]
                    logging.info(f'--> {header} extracted.')
                except Exception as e:
                    pass

            year_dropdown = driver.find_element(By.XPATH, "//select[@ng-model='periodoAnho']")
            year_select = Select(year_dropdown)
            year_select.select_by_value(year)
            time.sleep(3)

    def _registro_de_compras_y_venta(self, driver: webdriver.Chrome, doc):
        years = ["2023", "2024"]
        months = ["01", "02", "03", "04", "05", "06", "07", "08", "09", "10", "11", "12"]
        for year in years:
            for month in months:
                try:
                    month_dropdown = driver.find_element(By.ID, "periodoMes")
                    month_select = Select(month_dropdown)
                    month_select.select_by_value(month)
                    time.sleep(1)
                    year_dropdown = driver.find_element(By.XPATH, "//select[@ng-model='periodoAnho']")
                    year_select = Select(year_dropdown)
                    year_select.select_by_value(year)
                    time.sleep(1)
                    driver.find_element(By.XPATH, "//button[normalize-space()='Consultar']").click()
                    table_xpath = "//table[@class='table table-sm ng-scope']//tbody"
                    time.sleep(2)
                    driver.find_element(By.XPATH, "//a[@href='#venta/']").click()
                    WebDriverWait(driver, 5).until(EC.presence_of_element_located((By.XPATH, table_xpath)))
                    html = driver.find_elements(By.XPATH, table_xpath)[1].get_attribute('innerHTML')
                    soup = BeautifulSoup(html, 'html.parser')
                    td_elements = soup.find_all('td')
                    td_values = [td.get_text(strip=True) for td in td_elements]
                    sub_lists = [td_values[i:i + 8] for i in range(0, len(td_values), 8)]

                    header = driver.find_element(By.XPATH, "//strong[@class='ng-binding']").text
                    doc.add_heading(header, 1)
                    table = doc.add_table(rows=len(sub_lists) + 1, cols=8)
                    table.style = 'Table Grid'
                    table.cell(0, 0).text = 'Tipo Documento'
                    table.cell(0, 1).text = 'Total Documentos'
                    table.cell(0, 2).text = 'Monto Exento'
                    table.cell(0, 3).text = 'Monto Neto'
                    table.cell(0, 4).text = 'Monto IVA'
                    table.cell(0, 5).text = 'Monto Total'
                    for ind, line in enumerate(sub_lists):
                        table.cell(ind + 1, 0).text = line[0]
                        table.cell(ind + 1, 1).text = line[1]
                        table.cell(ind + 1, 2).text = line[2]
                        table.cell(ind + 1, 3).text = line[3]
                        table.cell(ind + 1, 4).text = line[4]
                        table.cell(ind + 1, 5).text = line[5]
                    logging.info(f'--> {header} extracted.')
                except Exception as e:
                    pass

            year_dropdown = driver.find_element(By.XPATH, "//select[@ng-model='periodoAnho']")
            year_select = Select(year_dropdown)
            year_select.select_by_value(year)
            time.sleep(3)

    def parse_observation_table(self, html):
        soup = BeautifulSoup(html, 'html.parser')

        tabla = soup.find("table", attrs={"border": "1"})

        if not tabla:
            return []

        data = []
        filas = tabla.find_all("tr")

        for fila in filas:
            celdas = fila.find_all("td")
            if len(celdas) >= 3:
                observacion = celdas[0].get_text(strip=True)
                descripcion = celdas[1].get_text(strip=True)
                disc_casos = celdas[2].get_text(strip=True)
                data.append({
                    "observacion": observacion,
                    "descripcion": descripcion,
                    "disc_casos": disc_casos
                })

        return data

    meses_map = {
        "enero": 1, "febrero": 2, "marzo": 3, "abril": 4,
        "mayo": 5, "junio": 6, "julio": 7, "agosto": 8,
        "septiembre": 9, "octubre": 10, "noviembre": 11, "diciembre": 12
    }

    def clave_orden(self, fecha_titulo):
        match = re.search(r"Observaciones: (\w+) de (\d{4})", fecha_titulo, re.IGNORECASE)
        if match:
            mes = match.group(1).lower()
            anio = int(match.group(2))
            mes_num = self.meses_map.get(mes, 0)
            return (anio, mes_num)
        return (0, 0)

    def _page1(self, driver: webdriver.Chrome, doc):
        # misii
        self._misii_doc(driver, doc)

        # click on `Datos personales y tributarios` menu and wait until dom change
        click_ele = driver.find_element(By.XPATH, "//a[@id='menu_datos_contribuyente']")
        driver.execute_script("arguments[0].click();", click_ele)
        self._wait_until_text_visible(driver, "//div[@id='box_datos_contribuyente']//h1[@class='title']",
                                      "Datos personales y tributarios")

        # Inicio de actividades y término de giro
        self._inicio_de_actividades_y_termino_de_giro(driver, doc)

        # Representantes legales vigentes
        self._representantes_legales_vigentes(driver, doc)

        # Socios y Capital
        self._socios_y_capital(driver, doc)

        # Actividades económicas
        self._actividades_economicas(driver, doc)

        # Sociedades a las que pertenece el contribuyente
        self._sociedades_a_las_que_pertenece_el_contribuyente(driver, doc)

        # Características del contribuyente
        self._caracteristicas_del_contribuyente(driver, doc)

        # Bienes Raíces
        self._bienes_raices(driver, doc)

        # click on `SII te informa` menu and wait until dom change
        click_ele = driver.find_element(By.XPATH, "//a[@id='menu_comunicados_sii']")
        driver.execute_script("arguments[0].click();", click_ele)
        self._wait_until_text_visible(driver, "//div[@id='box_comunicados_sii']//h1[contains(@class,'title')]",
                                      "SII te informa")

        # Anotaciones Vigentes
        self._anotaciones_vigentes(driver, doc)

        # Notificaciones
        self._notificaciones(driver, doc)

    def _page2(self, driver: webdriver.Chrome, doc):
        driver.get('https://www4.sii.cl/consdcvinternetui/#/index')
        time.sleep(5)

        # REGISTRO DE COMPRAS Y CAMPRA
        self._registro_de_compras_y_compra(driver, doc)

        # REGISTRO DE COMPRAS Y VENTA
        self._registro_de_compras_y_venta(driver, doc)

    def _page3(self, driver: webdriver.Chrome, doc):
        driver.get('https://www4.sii.cl/sifmConsultaInternet/index.html?form=29&dest=cifxx')
        time.sleep(5)
        try:
            f29 = "//a[normalize-space()='F29 (+)']"
            WebDriverWait(driver, 300).until(EC.presence_of_element_located((By.XPATH, f29)))
            driver.find_element(By.XPATH, f29).click()
            time.sleep(5)
            table_css = "body > div:nth-child(13) > div:nth-child(2) > div:nth-child(1) > table:nth-child(1) > tbody:nth-child(1) > tr:nth-child(2) > td:nth-child(1) > table:nth-child(1) > tbody:nth-child(1) > tr:nth-child(1) > td:nth-child(1) > div:nth-child(1) > table:nth-child(5) > tbody:nth-child(1) > tr:nth-child(1) > td:nth-child(2) > table:nth-child(1) > tbody:nth-child(2) > tr:nth-child(1) > td:nth-child(1) > table:nth-child(1) > tbody:nth-child(1) > tr:nth-child(1) > td:nth-child(1) > table:nth-child(1) > tbody:nth-child(2) > tr:nth-child(3) > td:nth-child(1) > table:nth-child(1) > tbody:nth-child(1) > tr:nth-child(2) > td:nth-child(1) > table:nth-child(1) > tbody:nth-child(1) > tr:nth-child(1) > td:nth-child(1) > table:nth-child(1) > tbody:nth-child(2) > tr:nth-child(1) > td:nth-child(2) > table:nth-child(1) > tbody:nth-child(2)"
            html = driver.find_element(By.CSS_SELECTOR, table_css).get_attribute('innerHTML')
            soup = BeautifulSoup(html, 'html.parser')
            td_elements = soup.find_all('td')
            observacion_indices = []
            for i, td in enumerate(td_elements):
                text = td.get_text(strip=True)
                if "SII" in text:
                    observacion_indices.append(i)
    

            tabla_elemento = driver.find_element(By.XPATH, '//*[@id="frame-window"]/table/tbody/tr[2]/td/table/tbody/tr/td/div/table[2]/tbody/tr/td[2]/table/tbody/tr[1]/td/table/tbody/tr[1]/td/table/tbody/tr[3]/td/table')
            driver.execute_script("arguments[0].scrollIntoView();", tabla_elemento)
            time.sleep(1)
            tabla_elemento.screenshot("tabla_general.png")

            observaciones = []
            selenium_tds = driver.find_elements(By.CSS_SELECTOR, table_css + " td")

            for idx in observacion_indices:
                try:
                    td_element = selenium_tds[idx]
                    img = td_element.find_element(By.CSS_SELECTOR, "img")
                    img.click()

                    WebDriverWait(driver, 10).until(
                        EC.presence_of_element_located((By.CSS_SELECTOR, "h2.gw-h2"))
                    )
                    WebDriverWait(driver, 10).until(
                        EC.presence_of_element_located((By.CSS_SELECTOR, "table[border='1']"))
                    )

                    obs_html = driver.find_element(By.CSS_SELECTOR, "table[border='1']").get_attribute('outerHTML')

                    titulo_periodo = driver.find_element(By.CSS_SELECTOR, "h2.gw-h2").text
                    match = re.search(r"Periodo\s+(\w+)\s+de\s+(\d{4})", titulo_periodo)
                    if match:
                        mes = match.group(1)
                        anio = match.group(2)
                        titulo = f"Observaciones: {mes} de {anio}"
                    else:
                        titulo = f"Observaciones {idx + 1}"

                    obs_data = self.parse_observation_table(obs_html)
                    observaciones.append({"titulo": titulo, "datos": obs_data})

                    driver.get('https://www4.sii.cl/sifmConsultaInternet/index.html?form=29&dest=cifxx')
                    WebDriverWait(driver, 30).until(
                        EC.presence_of_element_located((By.XPATH, "//a[normalize-space()='F29 (+)']"))
                    )
                    driver.find_element(By.XPATH, "//a[normalize-space()='F29 (+)']").click()
                    time.sleep(5)

                    selenium_tds = driver.find_elements(By.CSS_SELECTOR, table_css + " td")

                except Exception as e:
                    logging.warning(f"Error extrayendo observación en idx={idx}: {str(e)}")
                    continue

            doc.add_heading("CONSULTA INTEGRAL DE FISCALIZACION", 1)
            doc.add_picture("tabla_general.png", width=Inches(6))

            observaciones.sort(key=lambda x: self.clave_orden(x["titulo"]))

            for bloque in observaciones:
                titulo = bloque["titulo"]
                datos = bloque["datos"]

                if not datos:
                    continue

                doc.add_heading(titulo, level=2)
                table = doc.add_table(rows=1 + len(datos), cols=3)
                table.style = 'Table Grid'
                table.autofit = True

                # Encabezado
                table.cell(0, 0).text = 'Observación'
                table.cell(0, 1).text = 'Descripción'
                table.cell(0, 2).text = 'Disc./Nº Casos'

                for idx, obs in enumerate(datos):
                    table.cell(idx + 1, 0).text = obs["observacion"]
                    table.cell(idx + 1, 1).text = obs["descripcion"]
                    table.cell(idx + 1, 2).text = obs["disc_casos"]
            

            logging.info('--> CONSULTA INTEGRAL DE FISCALIZACION extracted.')
            self._fix_three_col_widths(table)

        except Exception as e:
            pass

    def _page4(self, driver: webdriver.Chrome, doc):
        year_value = ["2022", "2023", "2024"]
        month_value = ["Enero", "Febrero", "Marzo", "Abril", "Mayo", "Junio", "Julio", "Agosto", "Septiembre",
                       "Octubre", "Noviembre", "Diciembre"]
        given_codes = ['586', '142', '503', '502', '110', '511', '514', '564', '521', '566', '560', '584', '562', '519',
                       '520', '534', '535', '563', '48', '91']

        for year in year_value:
            for month in month_value:
                try:
                    codes = []
                    driver.get('https://www4.sii.cl/rfiInternet/consulta/index.html#rfiSelFormularioPeriodo')
                    driver.refresh()
                    select_box_title_xpath = "//button[normalize-space()='Buscar Datos Ingresados']"
                    WebDriverWait(driver, 30).until(EC.presence_of_element_located((By.XPATH, select_box_title_xpath)))
                    first_dropdown = driver.find_elements(By.XPATH, '//select[@class="gwt-ListBox"]')[0]
                    first_select = Select(first_dropdown)
                    first_select.select_by_value("0")
                    time.sleep(5)
                    second_dropdown = driver.find_elements(By.XPATH, '//select[@class="gwt-ListBox"]')[1]
                    second_select = Select(second_dropdown)
                    second_select.select_by_visible_text(year)
                    time.sleep(3)
                    third_dropdown = driver.find_elements(By.XPATH, '//select[@class="gwt-ListBox"]')[2]
                    third_select = Select(third_dropdown)
                    third_select.select_by_visible_text(month)
                    time.sleep(3)
                    title_ele = driver.find_element(By.XPATH, select_box_title_xpath)
                    driver.execute_script("arguments[0].click();", title_ele)
                    time.sleep(3)

                    result_xpath = "//td[normalize-space()='DECLARACIONES VIGENTES']"
                    WebDriverWait(driver, 30).until(EC.presence_of_element_located((By.XPATH, result_xpath)))

                    links_ele = driver.find_elements(By.XPATH, '//div[@class="gwt-Hyperlink"]//a')
                    if len(links_ele) > 2:
                        driver.find_elements(By.XPATH, '//div[@class="gwt-Hyperlink"]//a')[2].click()
                        time.sleep(5)
                        driver.find_element(By.XPATH, "//button[normalize-space()='Ver Datos']").click()
                        code_table_xpath = "//div[@class='gwt-HTML']//div[3]"
                        WebDriverWait(driver, 30).until(EC.presence_of_element_located((By.XPATH, code_table_xpath)))
                        trs = driver.find_elements(By.XPATH, "//div[@class='gwt-HTML']//div//tr")
                        for tr in trs:
                            tr_html = tr.get_attribute('innerHTML')
                            if tr_html:
                                soup = BeautifulSoup(tr_html, 'html.parser')
                                td_elements = soup.find_all('td')
                                td_values = [td.get_text(strip=True) for td in td_elements]
                                filtered_items = [item for item in td_values if item not in ['+', '-', '=']]
                                if any(code in filtered_items for code in given_codes) and len(filtered_items) == 6:
                                    codes.append(filtered_items)

                        doc.add_heading(f'FORMULARIO 29 - {month} - {year}', 1)
                        table = doc.add_table(rows=len(codes) + 1, cols=6)
                        table.style = 'Table Grid'
                        for ind, line in enumerate(codes):
                            table.cell(ind + 1, 0).text = line[0]
                            table.cell(ind + 1, 1).text = line[1]
                            table.cell(ind + 1, 2).text = line[2]
                            table.cell(ind + 1, 3).text = line[3]
                            table.cell(ind + 1, 4).text = line[4]
                            table.cell(ind + 1, 5).text = line[5]

                        logging.info(f"--> FORMULARIO 29 - {month} - {year}' extracted.")

                except Exception as e:
                    pass

    def _page5(self, driver: webdriver.Chrome, doc):
        year_value = ["2022", "2023", "2024"]

        for year in year_value:
            try:
                driver.get('https://www4.sii.cl/consultaestadof22ui/#!/default')
                dropdown_xpath = "//select[@class='form-control ng-pristine ng-untouched ng-valid ng-not-empty']"
                WebDriverWait(driver, 30).until(EC.presence_of_element_located((By.XPATH, dropdown_xpath)))
                time.sleep(5)
                dropdown = driver.find_element(By.XPATH, dropdown_xpath)
                dropdown_select = Select(dropdown)
                dropdown_select.select_by_visible_text(year)
                time.sleep(3)
                driver.find_element(By.XPATH, "//button[normalize-space()='Consultar']").click()
                time.sleep(3)
                social_path = "//div[@id='SituacionActual']//span"
                WebDriverWait(driver, 30).until(EC.presence_of_element_located((By.XPATH, social_path)))
                social = driver.find_element(By.XPATH, social_path).text
                time.sleep(10)
                ver_formulario_22_ompacto_button = "//button[normalize-space()='Ver Formulario 22 Compacto']"
                WebDriverWait(driver, 30).until(
                    EC.presence_of_element_located((By.XPATH, ver_formulario_22_ompacto_button)))
                driver.find_element(By.XPATH, ver_formulario_22_ompacto_button).click()

                codes = ["1412", "1430", "1729", "36", "1445", "1459", "1484", "1496", "1513", "1564", "91", "87"]
                data = []

                # 1st container
                table_xpath = '//div[@class="container-fluid container-card no-cabecera"]//div[@class="div-table"]'
                WebDriverWait(driver, 30).until(EC.presence_of_element_located((By.XPATH, table_xpath)))
                tables_ele = driver.find_elements(By.XPATH, table_xpath)
                for item in tables_ele:
                    lines = item.text.split(' ')
                    code = lines[0]
                    name = " ".join(lines[1:-1])
                    number = lines[-1]
                    if code in codes:
                        data.append([code, name, number])

                # 2nd container
                table_xpath = '//div[@class="container-fluid container-card container-result"]//div[@class="div-table"]'
                WebDriverWait(driver, 30).until(EC.presence_of_element_located((By.XPATH, table_xpath)))
                tables_ele = driver.find_elements(By.XPATH, table_xpath)
                for item in tables_ele:
                    lines = item.text.split(' ')
                    lines.pop(-1)
                    matching_indices = [index for index, value in enumerate(lines) if value in codes]
                    if matching_indices:
                        code = lines[matching_indices[0]]
                        lines.pop(matching_indices[0])
                        name = " ".join(lines[0:matching_indices[0]])
                        if len(lines) > matching_indices[0]:
                            number = lines[-1]
                        else:
                            number = ''
                        data.append([code, name, number])

                doc.add_heading(f'CONSULTA DE ESTADO DE DECLARACIÓN DE RENTA {year}', 1)
                doc.add_paragraph(f'Situación Renta Actual: {social}')
                doc.add_heading(f"AÑO TRIBUTARIO {year}", 1)
                # ━━━ build & fill the 3‑column table ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
                table = doc.add_table(rows=len(data) + 1, cols=3)
                table.style = 'Table Grid'
                table.cell(0, 0).text = 'Code'
                table.cell(0, 1).text = 'Name'
                table.cell(0, 2).text = 'Number'
                for ind, line in enumerate(data):
                    table.cell(ind + 1, 0).text = line[0]
                    table.cell(ind + 1, 1).text = line[1]
                    table.cell(ind + 1, 2).text = line[2]

                # ── NEW: force widths  (Code ≈1ʺ, Name ≈4ʺ, Number ≈1ʺ) ──────────
                from docx.shared import Inches                      # put at top of file once
                widths = [Inches(1), Inches(4), Inches(1)]          # tweak as you like
                for idx, w in enumerate(widths):
                    for cell in table.columns[idx].cells:
                        cell.width = w
                table.autofit = False                               # stop Word re‑stretching
                # ──────────────────────────────────────────────────────────────────


                logging.info(f'Ver Formulario 22 Compacto - {year} extracted')

            except Exception as e:
                pass

    def _page6(self, driver: webdriver.Chrome, doc):
        try:
            driver.get('https://www2.sii.cl/pagogiro-ui/home')
            link_xpath = "//a[normalize-space()='Consulta y Pago de Giros']"
            WebDriverWait(driver, 120).until(EC.presence_of_element_located((By.XPATH, link_xpath)))
            click_ele = driver.find_element(By.XPATH, link_xpath)
            driver.execute_script("arguments[0].click();", click_ele)
            table_xpath = "//table[@id='table-giros']"
            WebDriverWait(driver, 120).until(EC.presence_of_element_located((By.XPATH, table_xpath)))
            html = driver.find_element(By.XPATH, table_xpath).get_attribute('innerHTML')
            soup = BeautifulSoup(html, 'html.parser')
            td_elements = soup.find_all('td')
            td_values = [td.get_text(strip=True) for td in td_elements]
            sub_lists = [td_values[i:i + 9] for i in range(0, len(td_values), 9)]
            doc.add_heading('Consulta y Pago de Giros', 1)
            table = doc.add_table(rows=len(sub_lists) + 1, cols=8)
            table.style = 'Table Grid'
            table.cell(0, 0).text = 'Folio'
            table.cell(0, 1).text = 'Fec. Emisión'
            table.cell(0, 2).text = 'Fec. Vcto'
            table.cell(0, 3).text = 'Negocio Emisor'
            table.cell(0, 4).text = 'Moneda'
            table.cell(0, 5).text = 'Form'
            table.cell(0, 6).text = 'Total a Pagar'
            table.cell(0, 7).text = 'Acciones'
            for ind, line in enumerate(sub_lists):
                table.cell(ind + 1, 0).text = line[1]
                table.cell(ind + 1, 1).text = line[2]
                table.cell(ind + 1, 2).text = line[3]
                table.cell(ind + 1, 3).text = line[4]
                table.cell(ind + 1, 4).text = line[5]
                table.cell(ind + 1, 5).text = line[6]
                table.cell(ind + 1, 6).text = line[7]
                table.cell(ind + 1, 7).text = line[8]
            logging.info('--> Consulta y Pago de Giros extracted.')
        except Exception as e:
            pass

    def _page7(self, driver: webdriver.Chrome, doc):
        years = ["2023", "2024", "2022"]
        for year in years:
            try:
                driver.get('https://www4.sii.cl/perfilamientodjui/#/declaracionJuradaRenta')
                dropdown_xpath = "//select[@name='anioTributario']"
                WebDriverWait(driver, 50).until(EC.presence_of_element_located((By.XPATH, dropdown_xpath)))
                dropdown = driver.find_element(By.XPATH, dropdown_xpath)
                dropdown_select = Select(dropdown)
                dropdown_select.select_by_visible_text(year)
                time.sleep(5)
                table_xpath = "//table[@st-table='datosMostrarComunes']"
                WebDriverWait(driver, 30).until(EC.presence_of_element_located((By.XPATH, table_xpath)))
                html = driver.find_element(By.XPATH, table_xpath).get_attribute('innerHTML')
                soup = BeautifulSoup(html, 'html.parser')
                td_elements = soup.find_all('td')
                td_values = [td.get_text(strip=True) for td in td_elements]
                sub_lists = [td_values[i:i + 6] for i in range(0, len(td_values), 6)]
                doc.add_heading(f'Declaraciones por Régimen - {year}', 1)
                table = doc.add_table(rows=len(sub_lists) + 1, cols=4)
                table.style = 'Table Grid'
                table.cell(0, 0).text = 'codigo'
                table.cell(0, 1).text = 'Declaración Jurada'
                table.cell(0, 2).text = 'Fec. Vcto'
                table.cell(0, 3).text = 'Fecha Presentación'
                for ind, line in enumerate(sub_lists):
                    if len(line) > 4:
                        table.cell(ind + 1, 0).text = line[0]
                        table.cell(ind + 1, 1).text = line[1]
                        table.cell(ind + 1, 2).text = line[2]
                        table.cell(ind + 1, 3).text = line[3]
                    
            except Exception as e:
                pass

    def _page8(self, driver: webdriver.Chrome, doc):
        # Obtener la fecha actual
        fecha_actual = datetime.now()

        # Evaluar el mes actual y asignar el año correspondiente
        if fecha_actual.month >= 6:
            year = str(fecha_actual.year)
        else:
            year = str(fecha_actual.year - 1)
        try:
            driver.get('https://www4.sii.cl/consultaestadof22ui/#!/default')
            dropdown_xpath = "//*[@id=\"formulario-periodo\"]/div/div[2]/div/select"
            WebDriverWait(driver, 50).until(EC.presence_of_element_located((By.XPATH, dropdown_xpath)))
            dropdown = driver.find_element(By.XPATH, dropdown_xpath)
            dropdown_select = Select(dropdown)
            dropdown_select.select_by_visible_text(year)
            query_btn = "//*[@id=\"formulario-periodo\"]/div/div[2]/div/button"
            WebDriverWait(driver, 30).until(EC.presence_of_element_located((By.XPATH, query_btn)))
            driver.find_element(By.XPATH, query_btn).click()
            query_btn = "//*[@id=\"my-wrapper\"]/div[3]/div/div/div/div/div[2]/div[2]/div/table/tbody/tr[1]/td[3]/a"
            WebDriverWait(driver, 30).until(EC.presence_of_element_located((By.XPATH, query_btn)))
            driver.find_element(By.XPATH, query_btn).click()
            f22 = "encabezado"
            WebDriverWait(driver, 30).until(EC.presence_of_element_located((By.ID, f22)))
            time.sleep(2)
            
            codigos = [
                "1301", "1498", "1382", "1260", "1730", "1752", "1313", "1449", "1452",
                "1246", "1234", "1453", "1270", "1495", "1497", "1288", "1279", "1496",
                "1200", "1451", "1221", "1454", "1300", "1210"
            ]
            v = {}
            for codigo in codigos:
                try:
                    # Obtener el elemento por ID
                    print(codigo)
                    elemento = driver.find_element(By.ID, codigo)

                    # Forzar lectura del valor real usando JS (mejor para campos Angular)
                    valor = driver.execute_script("return arguments[0].value;", elemento)
                    if not valor:
                        valor = elemento.text

                    # Guardar usando clave tipo 'c1301', 'c1498', etc.
                    if valor:
                        v[f"{codigo}"] = int(valor.strip().replace(".", ""))
                    else:
                        v[f"{codigo}"] = 0

                except NoSuchElementException:
                    print(f"Elemento con ID {codigo} no encontrado.")
                    v[f"c{codigo}"] = 0  # Puedes asignar "" o 0 si prefieres
            # Factores
            factor_no_restitucion = 0.142857
            factor_restitucion = 0.369863

            # Cálculos de distribución
            no_restitucion = (v["1279"] + v["1496"]) / factor_no_restitucion if (v["1279"] + v["1496"]) > 0 else 0
            restitucion = (v["1300"] + v["1498"]) / factor_restitucion if (v["1300"] + v["1498"]) > 0 else 0

            # Valores base
            tabla_datos = [
                ["No sujeto a restitución", no_restitucion, str(v["1270"] + v["1495"]), str(v["1279"] + v["1496"]), factor_no_restitucion],
                ["Sujeto a restitución", restitucion, str(v["1497"] + v["1288"]), str(v["1300"] + v["1498"]), factor_restitucion],
                ["STUT", v["1382"] + v["1260"], "", "", ""],
                ["ISFUT", v["1730"] + v["1752"], "", "", ""],
                ["IPE", v["1313"] + v["1449"], "", "RAI", str(v["1210"] + v["1451"])],
                ["RAP", v["1452"] + v["1221"], "", "", ""],
                ["INR", v["1454"] + v["1246"], "", "", ""],
                ["RENTA EXENTA", v["1234"] + v["1453"], "", "", ""],
            ]

            print(v["1200"] + v["1451"])

            # Total
            total = sum(f[1] for f in tabla_datos if isinstance(f[1], (int, float)))
            tabla_datos.append(["Total", total, "", "", ""])

            # Documento
            doc.add_heading(f'Detalle de utilidades disponibles para distribución según RTRE al 31-12-{year}', 1)

            # Tabla con 5 columnas
            table = doc.add_table(rows=1, cols=5)
            table.style = 'Table Grid'

            # Encabezados
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Tipo de utilidades'
            hdr_cells[1].text = 'Cantidad a distribuir'
            hdr_cells[2].text = 'Crédito asociado\nSin D° Devolución'
            hdr_cells[3].text = 'Crédito asociado\nCon D° Devolución'
            hdr_cells[4].text = 'Factor'

            for row in tabla_datos:
                cells = table.add_row().cells
                
                # Columna 0: Tipo de utilidades
                cells[0].text = row[0]

                # Columna 1: Cantidad a distribuir (mostrar solo si > 0)
                if isinstance(row[1], (int, float)) and row[1] != 0:
                    cells[1].text = f"{row[1]:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")
                else:
                    cells[1].text = ""

                # Columna 2: Crédito asociado sin devolución
                if isinstance(row[2], (int, float)) and row[2] != 0:
                    cells[2].text = str(row[2])
                elif isinstance(row[2], str):
                    cells[2].text = row[2]
                else:
                    cells[2].text = ""

                # Columna 3: Crédito asociado con devolución
                if isinstance(row[3], (int, float)) and row[3] != 0:
                    cells[3].text = str(row[3])
                elif isinstance(row[3], str):
                    cells[3].text = row[3]
                else:
                    cells[3].text = ""

                # Columna 4: Factor
                if isinstance(row[4], float) and row[4] != 0:
                    cells[4].text = f"{row[4]:.6f}"
                elif isinstance(row[4], str):
                    cells[4].text = row[4]
                else:
                    cells[4].text = ""
                

                        
        except Exception as e:
            pass


    def run(self):
        logging.info('----------------- SCRIPT STARTS -------------------')
        driver = self.config_driver()
        # driver.minimize_window()
        doc = Document()
        try:
            self._login(driver)
            logging.info('==== PAGE 1 STARTED ====')
            self._page1(driver, doc)
            logging.info('==== PAGE 2 STARTED ====')
            #self._page2(driver, doc)
            logging.info('==== PAGE 3 STARTED ====')
            self._page3(driver, doc)
            logging.info('==== PAGE 4 STARTED ====')
            #self._page4(driver, doc)
            logging.info('==== PAGE 5 STARTED ====')
            self._page5(driver, doc)
            logging.info('==== PAGE 6 STARTED ====')
            self._page6(driver, doc)
            logging.info('==== PAGE 7 STARTED ====')
            self._page7(driver, doc)
            logging.info('==== PAGE 8 STARTED ====')
            self._page8(driver, doc)
            # Save the document in the specified output directory if provided
            if self.output_dir:
                timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
                output_path = os.path.join(
                    self.output_dir, f'{self.rut}_{timestamp}.docx')
                doc.save(output_path)
                logging.info(f'Document saved to: {output_path}')
                return [output_path]
            else:
                doc.save(f'{self.rut}.docx')
                logging.info(f'Document saved to: {self.rut}.docx')
        except Exception as e:
            logging.error(f"Error in script execution: {e}", exc_info=True)
        finally:
            # Ensure driver is properly closed
            try:
                driver.quit()
                logging.info("Browser driver closed successfully")
            except Exception as e:
                logging.error(f"Error closing driver: {e}", exc_info=True)
        logging.info('----------------- SCRIPT ENDS -------------------')
